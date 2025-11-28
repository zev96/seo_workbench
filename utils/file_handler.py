"""
文件处理工具
处理 Excel 和 Word 文件的读写操作
"""

import os
from pathlib import Path
from typing import List, Optional
import pandas as pd
from docx import Document
from loguru import logger


class FileHandler:
    """文件处理工具类"""
    
    @staticmethod
    def read_excel(file_path: str) -> Optional[pd.DataFrame]:
        """
        读取 Excel 文件
        
        Args:
            file_path: Excel 文件路径
            
        Returns:
            DataFrame 对象，失败返回 None
        """
        try:
            if not os.path.exists(file_path):
                logger.error(f"文件不存在: {file_path}")
                return None
            
            # 读取 Excel
            df = pd.read_excel(file_path, engine='openpyxl')
            logger.info(f"成功读取 Excel: {file_path}, 行数: {len(df)}, 列数: {len(df.columns)}")
            return df
            
        except Exception as e:
            logger.error(f"读取 Excel 失败: {e}")
            return None
    
    @staticmethod
    def write_excel(df: pd.DataFrame, file_path: str) -> bool:
        """
        写入 Excel 文件
        
        Args:
            df: DataFrame 对象
            file_path: 输出文件路径
            
        Returns:
            是否成功
        """
        try:
            # 确保目录存在
            Path(file_path).parent.mkdir(parents=True, exist_ok=True)
            
            # 写入 Excel
            df.to_excel(file_path, index=False, engine='openpyxl')
            logger.info(f"成功写入 Excel: {file_path}")
            return True
            
        except Exception as e:
            logger.error(f"写入 Excel 失败: {e}")
            return False
    
    @staticmethod
    def read_word(file_path: str) -> Optional[Document]:
        """
        读取 Word 文档
        
        Args:
            file_path: Word 文件路径
            
        Returns:
            Document 对象，失败返回 None
        """
        try:
            if not os.path.exists(file_path):
                logger.error(f"文件不存在: {file_path}")
                return None
            
            doc = Document(file_path)
            logger.info(f"成功读取 Word: {file_path}, 段落数: {len(doc.paragraphs)}")
            return doc
            
        except Exception as e:
            logger.error(f"读取 Word 失败: {e}")
            return None
    
    @staticmethod
    def create_word_from_template(template_path: str = None) -> Document:
        """
        创建 Word 文档（可基于模板）
        
        Args:
            template_path: 模板文件路径（可选）
            
        Returns:
            Document 对象
        """
        try:
            if template_path and os.path.exists(template_path):
                doc = Document(template_path)
                logger.info(f"基于模板创建 Word: {template_path}")
            else:
                doc = Document()
                logger.info("创建空白 Word 文档")
            return doc
            
        except Exception as e:
            logger.error(f"创建 Word 失败: {e}")
            return Document()
    
    @staticmethod
    def save_word(doc: Document, file_path: str) -> bool:
        """
        保存 Word 文档
        
        Args:
            doc: Document 对象
            file_path: 输出文件路径
            
        Returns:
            是否成功
        """
        try:
            # 确保目录存在
            Path(file_path).parent.mkdir(parents=True, exist_ok=True)
            
            # 保存文档
            doc.save(file_path)
            logger.info(f"成功保存 Word: {file_path}")
            return True
            
        except Exception as e:
            logger.error(f"保存 Word 失败: {e}")
            return False
    
    @staticmethod
    def get_image_files(folder_path: str, extensions: List[str] = None) -> List[Path]:
        """
        获取文件夹中的图片文件
        
        Args:
            folder_path: 文件夹路径
            extensions: 图片扩展名列表（默认: ['.jpg', '.jpeg', '.png', '.gif', '.bmp']）
            
        Returns:
            图片文件路径列表
        """
        if extensions is None:
            extensions = ['.jpg', '.jpeg', '.png', '.gif', '.bmp']
        
        try:
            folder = Path(folder_path)
            if not folder.exists():
                logger.warning(f"文件夹不存在: {folder_path}")
                return []
            
            # 收集所有图片文件
            image_files = []
            for ext in extensions:
                image_files.extend(folder.glob(f'*{ext}'))
                image_files.extend(folder.glob(f'*{ext.upper()}'))
            
            # 排序
            image_files = sorted(set(image_files))
            logger.info(f"找到 {len(image_files)} 个图片文件: {folder_path}")
            return image_files
            
        except Exception as e:
            logger.error(f"获取图片文件失败: {e}")
            return []
    
    @staticmethod
    def ensure_directory(dir_path: str) -> bool:
        """
        确保目录存在
        
        Args:
            dir_path: 目录路径
            
        Returns:
            是否成功
        """
        try:
            Path(dir_path).mkdir(parents=True, exist_ok=True)
            return True
        except Exception as e:
            logger.error(f"创建目录失败: {e}")
            return False
    
    @staticmethod
    def get_file_name_without_ext(file_path: str) -> str:
        """
        获取不带扩展名的文件名
        
        Args:
            file_path: 文件路径
            
        Returns:
            文件名（不含扩展名）
        """
        return Path(file_path).stem
    
    @staticmethod
    def list_word_files(folder_path: str) -> List[Path]:
        """
        列出文件夹中的所有 Word 文件
        
        Args:
            folder_path: 文件夹路径
            
        Returns:
            Word 文件路径列表
        """
        try:
            folder = Path(folder_path)
            if not folder.exists():
                logger.warning(f"文件夹不存在: {folder_path}")
                return []
            
            word_files = list(folder.glob('*.docx'))
            # 排除临时文件
            word_files = [f for f in word_files if not f.name.startswith('~$')]
            
            logger.info(f"找到 {len(word_files)} 个 Word 文件: {folder_path}")
            return sorted(word_files)
            
        except Exception as e:
            logger.error(f"列出 Word 文件失败: {e}")
            return []


if __name__ == "__main__":
    # 测试代码
    from .logger import setup_logger
    setup_logger()
    
    # 测试 Excel 读写
    test_data = pd.DataFrame({
        '标题': ['标题1', '标题2', '标题3'],
        '内容': ['内容1', '内容2', '内容3']
    })
    
    FileHandler.write_excel(test_data, 'test_output/test.xlsx')
    df = FileHandler.read_excel('test_output/test.xlsx')
    print(df)
    
    # 测试 Word 创建
    doc = FileHandler.create_word_from_template()
    doc.add_heading('测试标题', level=1)
    doc.add_paragraph('这是测试内容')
    FileHandler.save_word(doc, 'test_output/test.docx')

