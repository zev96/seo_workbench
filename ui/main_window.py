"""
主窗口
SEO 工作台的主界面，采用 Fluent Design 风格
"""

from PyQt6.QtWidgets import QWidget, QVBoxLayout, QHBoxLayout, QSplitter
from PyQt6.QtCore import Qt, QSize
from PyQt6.QtGui import QIcon
from qfluentwidgets import (
    MSFluentWindow, SplitFluentWindow, NavigationItemPosition,
    FluentIcon as FIF, SplashScreen
)
from loguru import logger

from .widgets.material_library import MaterialLibrary
from .widgets.smart_grid import SmartGrid
from .widgets.strategy_panel import StrategyPanel
from .widgets.toolbar import Toolbar
from .widgets.log_panel import LogPanel
from .widgets.comparison_table import ComparisonTableWidget
from .widgets.zhihu_monitor import ZhihuMonitorWidget
from .dialogs.api_settings import APISettingsDialog
from .dialogs.ai_title_dialog import AITitleDialog
from .dialogs.ai_rewrite_dialog import AIRewriteDialog
from .dialogs.strategy_config_dialog import StrategyConfigDialog
from .dialogs.seo_setting_dialog import SEOSettingDialog
from ..config.settings import ProfileConfig
from ..database.db_manager import DatabaseManager


class MainWindow(MSFluentWindow):
    """主窗口类 (Fluent 风格)"""
    
    def __init__(self, config: ProfileConfig, db_manager: DatabaseManager):
        super().__init__()
        
        self.config = config
        self.db_manager = db_manager
        
        # AI 标题队列（用于标题驱动混排）
        self.ai_title_queue = []
        self.ai_title_format = "H1"
        
        self._init_window()
        self._init_ui()
        self._connect_signals()
        
        logger.info("Fluent 主窗口初始化完成")
    
    def _init_window(self):
        """初始化窗口属性"""
        self.setWindowTitle("SEO 智能内容工作台")
        self.resize(1400, 900)
        
        # 设置窗口图标（这里暂时使用默认）
        # self.setWindowIcon(QIcon('resources/icons/app.ico'))
        
        # 开启 Mica 效果（Windows 11 毛玻璃背景）
        try:
            from qfluentwidgets import setTheme, Theme, setThemeColor
            from PyQt6.QtGui import QColor
            
            # 设置主题颜色为深天蓝
            setThemeColor(QColor("#4784d1"))
            
            # 开启 Mica 效果
            self.setMicaEffectEnabled(True)
            
            logger.info("Mica 毛玻璃效果已启用")
        except Exception as e:
            logger.warning(f"无法启用 Mica 效果: {e}")
        
        # 居中显示
        desktop = self.screen().availableGeometry()
        w, h = desktop.width(), desktop.height()
        self.move(w//2 - self.width()//2, h//2 - self.height()//2)
    
    def _init_ui(self):
        """初始化 UI"""
        # 创建中心部件容器（工作台）
        self.central_widget = QWidget()
        self.central_widget.setObjectName("central_widget")
        self.addSubInterface(self.central_widget, FIF.HOME, "工作台")
        
        # 主布局 (垂直：工具栏 + 内容 + 日志)
        # 增加 margins 营造呼吸感
        main_layout = QVBoxLayout(self.central_widget)
        main_layout.setContentsMargins(20, 20, 20, 20)
        main_layout.setSpacing(15)
        
        # 1. 顶部工具栏
        self.toolbar = Toolbar(self.config)
        main_layout.addWidget(self.toolbar)
        
        # 2. 中间内容区（三栏布局）
        # 使用 QSplitter 实现可调节的三栏
        content_splitter = QSplitter(Qt.Orientation.Horizontal)
        content_splitter.setHandleWidth(5)
        content_splitter.setStyleSheet("""
            QSplitter::handle {
                background-color: #e0e0e0;
                border-radius: 2px;
            }
            QSplitter::handle:hover {
                background-color: #009faa;
            }
        """)
        
        # 左侧：素材库
        self.material_library = MaterialLibrary(self.db_manager)
        content_splitter.addWidget(self.material_library)
        
        # 中间：网格编辑器
        self.smart_grid = SmartGrid(self.config)
        content_splitter.addWidget(self.smart_grid)
        
        # 右侧：策略面板
        self.strategy_panel = StrategyPanel(self.config)
        self.strategy_panel.set_validator(self._validate_strategy_columns)
        content_splitter.addWidget(self.strategy_panel)
        
        # 设置分割器比例（左:中:右 = 2:5:2）
        content_splitter.setStretchFactor(0, 20)
        content_splitter.setStretchFactor(1, 60)
        content_splitter.setStretchFactor(2, 20)
        
        main_layout.addWidget(content_splitter)
        
        # 创建对比表管理界面（新的子界面）
        self.comparison_table_widget = ComparisonTableWidget()
        self.addSubInterface(self.comparison_table_widget, FIF.DICTIONARY, "数据库")
        
        # 创建知乎监测界面
        self.zhihu_monitor_widget = ZhihuMonitorWidget(self.db_manager.get_session())
        self.zhihu_monitor_widget.setObjectName("zhihu_monitor_widget")
        self.addSubInterface(self.zhihu_monitor_widget, FIF.SEARCH, "知乎监测")
        
        # 3. 底部：日志面板（隐藏，用户不需要）
        # self.log_panel = LogPanel()
        # main_layout.addWidget(self.log_panel)
        
        # 添加其他导航项（设置页）
        # self.navigationInterface.addSeparator()
        
        self.navigationInterface.addItem(
            routeKey='settings',
            icon=FIF.SETTING,
            text='设置',
            onClick=self._on_api_settings,
            position=NavigationItemPosition.BOTTOM
        )
        
        self.navigationInterface.addItem(
            routeKey='about',
            icon=FIF.INFO,
            text='关于',
            onClick=self._on_about,
            position=NavigationItemPosition.BOTTOM
        )
    
    def _connect_signals(self):
        """连接信号和槽"""
        # 工具栏信号
        self.toolbar.generate_clicked.connect(self._on_generate)
        
        # 策略面板信号（右下角按钮）
        self.strategy_panel.import_excel_clicked.connect(self._on_import_excel)
        self.strategy_panel.clear_grid_clicked.connect(self.smart_grid.clear_all)
        self.strategy_panel.bold_tool_clicked.connect(self._on_bold_tool)
        self.strategy_panel.ai_title_clicked.connect(self._on_ai_title_dialog)
        self.strategy_panel.ai_rewrite_clicked.connect(self._on_ai_rewrite_dialog)
        self.strategy_panel.strategy_config_clicked.connect(self._on_strategy_config)
        self.strategy_panel.numbering_group_clicked.connect(self._on_numbering_group_config)
        self.strategy_panel.seo_config_clicked.connect(self._on_seo_config)
        self.strategy_panel.dedup_config_clicked.connect(self._on_dedup_config)
        
        # 素材库信号
        self.material_library.material_selected.connect(self._on_material_selected)
        
        # 网格信号
        self.smart_grid.data_changed.connect(self._on_grid_data_changed)
        
        # 初始化按钮状态（初始无数据，禁用相关按钮）
        self._update_button_states()
    
    def _validate_strategy_columns(self, columns: list[int]) -> tuple[bool, str]:
        """验证策略列号是否合法"""
        # 检查是否有数据
        active_rows = self.smart_grid.get_active_row_count()
        if active_rows == 0 and self.smart_grid.table.rowCount() == 0:
             # 如果连空表格都没有，或者没有任何内容
             # 实际上 SmartGrid 初始化时会创建表格，但可能是隐藏的
             # 我们认为如果 active_rows 为 0 且表格隐藏，则不允许设置
             if not self.smart_grid.table.isVisible():
                 return False, "工作区为空，请先导入数据或添加内容"

        # 检查列号是否越界
        max_col = self.smart_grid.table.columnCount()
        for col in columns:
            if col > max_col:
                return False, f"列号 {col} 超出当前工作区范围 (最大 {max_col} 列)"
        
        return True, ""

    # ==================== 槽函数 ====================
    
    def _on_import_excel(self):
        """导入 Excel"""
        self.smart_grid.import_from_excel()
        # 更新按钮状态（导入后有数据了）
        self._update_button_states()
    
    def _on_api_settings(self):
        """API 设置"""
        dialog = APISettingsDialog(self.config, self)
        if dialog.exec():
            logger.info("API 设置已更新")
    
    def _on_bold_tool(self):
        """批量加粗工具"""
        from .dialogs.bold_tool import BoldToolDialog
        dialog = BoldToolDialog(self.config, self)
        dialog.exec()
    
    def _on_about(self):
        """关于"""
        from qfluentwidgets import MessageBox
        w = MessageBox(
            "关于 SEO 工作台",
            "SEO 智能内容工作台 v5.0\n\n"
            "一个专为 SEO 团队设计的内容生产工具\n\n"
            "By CEWEY",
            self
        )
        w.yesButton.setText("确定")
        w.cancelButton.hide()
        w.exec()
    
    def _on_generate(self, mode: str):
        """生成文档"""
        logger.info(f"开始生成文档: 模式={mode}")
        from qfluentwidgets import InfoBar, InfoBarPosition
        from PyQt6.QtWidgets import QFileDialog
        from ..core.generation_worker import GenerationWorker
        from .dialogs.progress_dialog import ProgressDialog
        
        # 获取工作区数据
        grid_data = self.smart_grid.get_grid_data()
        if not grid_data:
            InfoBar.warning(
                title='提示',
                content='工作区为空，请先导入数据或添加内容',
                orient=Qt.Orientation.Horizontal,
                isClosable=True,
                position=InfoBarPosition.BOTTOM_RIGHT,
                duration=3000,
                parent=self
            )
            return
        
        # 选择保存路径
        save_dir = QFileDialog.getExistingDirectory(
            self,
            "选择保存目录",
            "",
            QFileDialog.Option.ShowDirsOnly
        )
        
        if not save_dir:
            return
        
        # 获取生成数量
        count = self.toolbar.count_spin.value() if mode == "shuffle" else len(grid_data)
        
        # 如果使用AI标题模式，确保数量与标题数量一致
        if len(self.ai_title_queue) > 0 and mode == "shuffle":
            count = len(self.ai_title_queue)
            logger.info(f"AI标题模式：强制使用标题数量 {count}")
        
        # 创建进度对话框
        progress_dialog = ProgressDialog(
            title="正在生成文档",
            total=count,
            parent=self
        )
        
        # 创建工作线程
        self.generation_worker = GenerationWorker(
            grid_data=grid_data,
            save_dir=save_dir,
            mode=mode,
            count=count,
            config=self.config,
            generate_func=self._generate_documents_with_progress,
            parent=self
        )
        
        # 连接信号
        self.generation_worker.progress_updated.connect(
            lambda current, total, detail: (
                progress_dialog.set_progress(current, total),
                progress_dialog.set_detail(detail)
            )
        )
        self.generation_worker.status_changed.connect(progress_dialog.set_status)
        self.generation_worker.generation_complete.connect(
            lambda success, msg, count: self._on_generation_complete(
                progress_dialog, success, msg, count, save_dir
            )
        )
        self.generation_worker.error_occurred.connect(
            lambda error: logger.error(f"生成错误: {error}")
        )
        
        # 连接取消信号
        progress_dialog.cancelled.connect(self.generation_worker.cancel)
        
        # 启动线程
        self.generation_worker.start()
        
        # 显示进度对话框
        progress_dialog.exec()
    
    def _on_generation_complete(self, dialog, success: bool, message: str, count: int, save_dir: str):
        """
        生成完成回调
        
        Args:
            dialog: 进度对话框
            success: 是否成功
            message: 消息
            count: 生成数量
            save_dir: 保存目录
        """
        from qfluentwidgets import InfoBar, InfoBarPosition
        
        # 更新对话框
        dialog.complete(success, message)
        
        # 显示通知
        if success:
            InfoBar.success(
                title='生成完成',
                content=f'已生成 {count} 个文档到 {save_dir}',
                orient=Qt.Orientation.Horizontal,
                isClosable=True,
                position=InfoBarPosition.BOTTOM_RIGHT,
                duration=5000,
                parent=self
            )
        else:
            InfoBar.error(
                title='生成失败',
                content=message,
                orient=Qt.Orientation.Horizontal,
                isClosable=True,
                position=InfoBarPosition.BOTTOM_RIGHT,
                duration=5000,
                parent=self
            )
        
        logger.info(f"生成完成: success={success}, count={count}")
    
    def _generate_documents_with_progress(
        self,
        grid_data: list,
        save_dir: str,
        mode: str,
        count: int,
        progress_callback=None
    ) -> int:
        """
        生成文档（支持进度回调）
        
        Args:
            grid_data: 网格数据
            save_dir: 保存目录
            mode: 生成模式
            count: 生成数量
            progress_callback: 进度回调函数 (current, total, detail)
        
        Returns:
            生成的文档数量
        """
        # 调用原有的生成方法，但添加进度回调
        return self._generate_documents(
            grid_data=grid_data,
            save_dir=save_dir,
            mode=mode,
            count=count,
            progress_callback=progress_callback
        )
    
    def _generate_documents(self, grid_data: list, save_dir: str, mode: str, count: int, progress_callback=None) -> int:
        """实际生成文档的逻辑"""
        import random
        from pathlib import Path
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
        from datetime import datetime
        from ..core.quality_checker import QualityChecker, QualityReport
        from ..core.smart_numbering import SmartNumbering
        
        generated = 0
        
        # 初始化质量检查器和报告
        quality_checker = None
        quality_report = None
        
        if self.config.quality_check_enabled:
            quality_checker = QualityChecker(
                threshold_premium=self.config.quality_threshold_premium,
                threshold_standard=self.config.quality_threshold_standard,
                seo_keywords=self.config.target_keywords if self.config.seo_check_enabled else [],
                seo_density_min=self.config.seo_density_min,
                seo_density_max=self.config.seo_density_max
            )
            quality_report = QualityReport()
            logger.info("质量检查已启用")
            if self.config.seo_check_enabled and self.config.target_keywords:
                logger.info(f"SEO 密度检查已启用，目标关键词: {self.config.target_keywords}")
        
        if mode == "row":
            # 按行生成模式：每行生成一个文档
            for idx, row_data in enumerate(grid_data):
                # 更新进度
                if progress_callback:
                    progress_callback(
                        idx + 1,
                        len(grid_data),
                        f"正在生成第 {idx + 1} 个文档..."
                    )
                
                doc = Document()
                
                # === 智能序号处理：按格式分类计数 ===
                # 为每种序号格式维护独立的计数器
                style_counters = {}
                
                # 根据列类型设置样式
                for col_idx, content in enumerate(row_data):
                    if not content or not content.strip():
                        continue
                    
                    col_type = self.config.get_column_type(col_idx)
                    
                    if col_type == 'Ignore':
                        continue
                    
                    # === 智能序号处理 ===
                    # 将内容按换行符分割成多个段落
                    paragraphs = content.split('\n')
                    
                    # 处理该列的每个段落
                    for para_text in paragraphs:
                        if not para_text.strip():
                            continue
                        
                        processed_content = para_text
                        
                        # === 对所有类型应用智能序号处理 ===
                        # 先检测是否有序号
                        cleaned_text, detected_style = SmartNumbering.detect_and_clean(para_text)
                        
                        if detected_style:
                            # 检测到序号：清洗并重新编号
                            # 为该格式初始化计数器（如果还没有）
                            if detected_style not in style_counters:
                                style_counters[detected_style] = 1
                            
                            current_number = style_counters[detected_style]
                            processed_content = SmartNumbering.process_text(
                                para_text,
                                current_number,
                                should_renumber=True
                            )
                            logger.info(f"[{col_type}] 重编号: {current_number}, 样式={detected_style}, 原文='{para_text[:40]}', 结果='{processed_content[:40]}'")
                            
                            # 递增该格式的计数器
                            style_counters[detected_style] += 1
                        else:
                            # 没有检测到序号：保持原样
                            processed_content = para_text
                            logger.debug(f"[{col_type}] 无序号，保持原样: '{para_text[:40]}'")
                        
                        # 根据类型添加段落
                        if col_type == 'H1':
                            p = doc.add_paragraph(processed_content)
                            self._apply_heading_style(p, level=1)
                        elif col_type == 'H2':
                            p = doc.add_paragraph(processed_content)
                            self._apply_heading_style(p, level=2)
                        elif col_type == 'H3':
                            p = doc.add_paragraph(processed_content)
                            self._apply_heading_style(p, level=3)
                        elif col_type == 'H4':
                            p = doc.add_paragraph(processed_content)
                            self._apply_heading_style(p, level=4)
                        elif col_type == 'List':
                            p = doc.add_paragraph(processed_content, style='List Bullet')
                            self._apply_body_style(p)
                        elif col_type == 'Body':
                            p = doc.add_paragraph(processed_content)
                            self._apply_body_style(p)
                        
                        # 应用加粗关键词
                        if col_type in ['Body', 'List'] and self.config.bold_keywords:
                            self._apply_bold_keywords(p, self.config.bold_keywords)
                    
                    # 插入该列的图片（如果有）- 在该列所有段落之后
                    self._insert_column_image(doc, col_idx)
                    
                    # 检查是否需要插入对比表图片（根据模式使用不同的变量名）
                    current_row_data = row_data if mode == "row" else processed_row
                    self._check_and_insert_comparison_table(doc, col_idx, content, current_row_data)
                
                # 质量检查和文件名标记
                title = row_data[0] if row_data else f"文档{idx + 1}"
                quality_prefix = ""
                
                if quality_checker:
                    # 创建文档指纹
                    fingerprint = quality_checker.create_fingerprint(row_data)
                    # 提取完整文本用于 SEO 检查
                    full_text = "\n".join([str(content) for content in row_data if content])
                    # 检查质量
                    score = quality_checker.check_quality(fingerprint, full_text)
                    # 添加前缀
                    quality_prefix = f"[{score.rating}]_"
                    # 记录到报告
                    if quality_report:
                        quality_report.add_record(
                            filename=f"{quality_prefix}文档_{idx + 1:04d}.docx",
                            title=title[:50],  # 限制长度
                            max_similarity=score.max_similarity,
                            rating=score.rating,
                            timestamp=datetime.now(),
                            keyword_density=score.keyword_density,
                            density_rating=score.density_rating,
                            seo_suggestion=score.seo_suggestion
                        )
                
                # 保存文档
                filename = f"{quality_prefix}文档_{idx + 1:04d}.docx"
                filepath = Path(save_dir) / filename
                doc.save(str(filepath))
                generated += 1
                
                logger.info(f"已生成文档 {generated}/{len(grid_data)}: {filename}")
                
        else:
            # 随机混排模式：应用混排策略
            # 检查是否启用了标题驱动模式
            use_ai_titles = len(self.ai_title_queue) > 0
            
            # 获取列数据（直接从表格按列获取，避免转置问题）
            columns_data = self.smart_grid.get_column_data()
            logger.info(f"获取列数据：共 {len(columns_data)} 列")
            for idx, col_data in enumerate(columns_data):
                logger.debug(f"列 {idx + 1}: {len(col_data)} 个有效内容")
            
            for i in range(count):
                # 更新进度
                if progress_callback:
                    progress_callback(
                        i + 1,
                        count,
                        f"正在生成第 {i + 1} 个文档（{'AI标题' if use_ai_titles else '混排'}模式）..."
                    )
                
                doc = Document()
                
                # 从每列独立随机选择内容（修复不等长列问题）
                processed_row = []
                for col_data in columns_data:
                    if col_data:
                        # 该列有内容，随机选择一个
                        content = random.choice(col_data)
                        processed_row.append(content)
                    else:
                        # 该列为空
                        processed_row.append("")
                
                # 应用混排策略（删除某些列）
                if self.config.shuffling_strategies:
                    processed_row = self._apply_column_shuffling_strategies(processed_row)
                
                # 标题驱动逻辑：如果有 AI 标题队列，替换第一列内容
                if use_ai_titles and i < len(self.ai_title_queue):
                    ai_title = self.ai_title_queue[i]
                    # 将 AI 标题插入到第一列
                    if len(processed_row) > 0:
                        processed_row[0] = ai_title
                    else:
                        processed_row = [ai_title]
                    
                    # 设置第一列的格式为 AI 指定的格式
                    self.config.set_column_type(0, self.ai_title_format, "AI标题")
                    logger.info(f"文档 {i+1}: 使用 AI 标题 '{ai_title}' (格式: {self.ai_title_format})")
                
                # === 智能序号处理：在混排策略之后，写入Word之前 ===
                # 按序号分组区间进行独立计数
                
                # 1. 创建列到序号分组的映射
                column_to_numbering_group = {}  # {列索引: 分组索引}
                
                if self.config.numbering_groups:
                    # 使用用户配置的序号分组
                    for group_idx, group_columns in enumerate(self.config.numbering_groups):
                        for col in group_columns:
                            column_to_numbering_group[col] = group_idx
                    logger.debug(f"使用序号分组配置: {self.config.numbering_groups}")
                else:
                    # 如果没有配置序号分组，则使用混排策略作为分组依据（兼容旧逻辑）
                    for strategy_idx, strategy in enumerate(self.config.shuffling_strategies):
                        for col in strategy.columns:
                            column_to_numbering_group[col] = strategy_idx
                    logger.debug(f"使用混排策略作为序号分组: {column_to_numbering_group}")
                
                # 2. 为每个分组维护独立的计数器字典
                group_counters = {}  # {分组索引: {序号样式: 计数器}}
                
                logger.debug(f"序号分组映射: {column_to_numbering_group}")
                
                # 处理每一列的内容（包括空列）
                for col_idx in range(len(processed_row)):
                    content = processed_row[col_idx] if col_idx < len(processed_row) else ""
                    
                    # 如果列有内容，处理内容
                    if content and content.strip():
                        col_type = self.config.get_column_type(col_idx)
                        
                        if col_type == 'Ignore':
                            # 忽略该列，跳过内容处理，但仍要检查对比表格
                            pass
                        else:
                            # === 智能序号处理 ===
                            # 将内容按换行符分割成多个段落
                            paragraphs = content.split('\n')
                            
                            # 处理该列的每个段落
                            for para_text in paragraphs:
                                if not para_text.strip():
                                    continue
                                
                                processed_content = para_text
                                
                                # === 对所有类型应用智能序号处理 ===
                                # 先检测是否有序号
                                cleaned_text, detected_style = SmartNumbering.detect_and_clean(para_text)
                                
                                if detected_style:
                                    # 检测到序号：判断是否需要重新编号
                                    
                                    # 确定该列属于哪个序号分组（注意：col_idx是代码索引，从0开始）
                                    group_idx = column_to_numbering_group.get(col_idx, -1)  # -1 表示不属于任何分组
                                    
                                    # 如果不在任何序号分组内，保持原样不重新编号
                                    if group_idx == -1:
                                        processed_content = para_text  # 保持原序号
                                        logger.debug(f"[{col_type}][列{col_idx+1}] 不在序号分组内，保持原样: '{para_text[:40]}'")
                                    else:
                                        # 在序号分组内，强制重新编号
                                        if group_idx not in group_counters:
                                            group_counters[group_idx] = {}
                                        current_counters = group_counters[group_idx]
                                        group_name = f"分组{group_idx+1}"
                                        
                                        # 为该格式初始化计数器（如果还没有）
                                        if detected_style not in current_counters:
                                            current_counters[detected_style] = 1
                                        
                                        current_number = current_counters[detected_style]
                                        
                                        # 强制使用计数器值重新生成序号前缀（修复原序号为1时的问题）
                                        new_prefix = SmartNumbering.generate_prefix(current_number, detected_style)
                                        processed_content = new_prefix + cleaned_text
                                        
                                        logger.info(f"[{col_type}][列{col_idx+1}][{group_name}] 重编号: {current_number}, 样式={detected_style}, 原文='{para_text[:40]}', 结果='{processed_content[:40]}'")
                                        
                                        # 递增该组该格式的计数器
                                        current_counters[detected_style] += 1
                                else:
                                    # 没有检测到序号：保持原样
                                    processed_content = para_text
                                    logger.debug(f"[{col_type}][列{col_idx}] 无序号，保持原样: '{para_text[:40]}'")
                                
                                # 根据类型添加段落
                                if col_type == 'H1':
                                    p = doc.add_paragraph(processed_content)
                                    self._apply_heading_style(p, level=1)
                                elif col_type == 'H2':
                                    p = doc.add_paragraph(processed_content)
                                    self._apply_heading_style(p, level=2)
                                elif col_type == 'H3':
                                    p = doc.add_paragraph(processed_content)
                                    self._apply_heading_style(p, level=3)
                                elif col_type == 'H4':
                                    p = doc.add_paragraph(processed_content)
                                    self._apply_heading_style(p, level=4)
                                elif col_type == 'List':
                                    p = doc.add_paragraph(processed_content, style='List Bullet')
                                    self._apply_body_style(p)
                                elif col_type == 'Body':
                                    p = doc.add_paragraph(processed_content)
                                    self._apply_body_style(p)
                                
                                # 应用加粗关键词
                                if col_type in ['Body', 'List'] and self.config.bold_keywords:
                                    self._apply_bold_keywords(p, self.config.bold_keywords)
                            
                            # 插入该列的图片（如果有）- 在该列所有段落之后
                            self._insert_column_image(doc, col_idx)
                    
                    # 立即检查该列的对比表格（无论列是否为空）
                    self._check_and_insert_comparison_table(doc, col_idx, content, processed_row)
                
                # 质量检查和文件名标记
                title = processed_row[0] if processed_row else f"文档{i + 1}"
                quality_prefix = ""
                
                if quality_checker:
                    # 创建文档指纹
                    fingerprint = quality_checker.create_fingerprint(processed_row)
                    # 提取完整文本用于 SEO 检查
                    full_text = "\n".join([str(content) for content in processed_row if content])
                    # 检查质量
                    score = quality_checker.check_quality(fingerprint, full_text)
                    # 添加前缀
                    quality_prefix = f"[{score.rating}]_"
                    # 记录到报告
                    if quality_report:
                        quality_report.add_record(
                            filename=f"{quality_prefix}{'AI标题文档' if use_ai_titles else '混排文档'}_{i + 1:04d}.docx",
                            title=title[:50],  # 限制长度
                            max_similarity=score.max_similarity,
                            rating=score.rating,
                            timestamp=datetime.now(),
                            keyword_density=score.keyword_density,
                            density_rating=score.density_rating,
                            seo_suggestion=score.seo_suggestion
                        )
                
                # 保存文档
                if use_ai_titles:
                    filename = f"{quality_prefix}AI标题文档_{i + 1:04d}.docx"
                else:
                    filename = f"{quality_prefix}混排文档_{i + 1:04d}.docx"
                filepath = Path(save_dir) / filename
                doc.save(str(filepath))
                generated += 1
                
                logger.info(f"已生成文档 {generated}/{count}: {filename}")
            
            # 生成完成后清空标题队列并解锁数量输入框
            if use_ai_titles:
                self.ai_title_queue = []
                # 🔓 解锁生成数量输入框
                self.toolbar.count_spin.setEnabled(True)
                self.toolbar.count_spin.setToolTip("")
                logger.info("AI 标题队列已清空，生成数量输入框已解锁")
        
        # 生成质量报告
        if quality_report and self.config.quality_generate_report:
            report_path = Path(save_dir) / "quality_report.csv"
            quality_report.save_to_csv(str(report_path))
            
            # 统计信息
            stats = quality_report.get_statistics()
            logger.info(f"查重统计: 优质={stats['查重_优质']}, 中等={stats['查重_中等']}, 高重复={stats['查重_高重复']}")
            if self.config.seo_check_enabled and self.config.target_keywords:
                logger.info(f"SEO统计: 完美={stats['SEO_完美']}, 不足={stats['SEO_不足']}, 堆砌={stats['SEO_堆砌']}")
        
        return generated
    
    def _insert_column_image(self, doc, col_idx: int):
        """为指定列插入随机图片
        
        Args:
            doc: Document对象
            col_idx: 列索引
        """
        import random
        from pathlib import Path
        from docx.shared import Inches, Cm
        from PIL import Image
        
        # 检查该列是否有图片组
        if col_idx not in self.config.column_images:
            return
        
        image_paths = self.config.column_images[col_idx]
        if not image_paths:
            return
        
        # 随机选择一张图片
        img_path = random.choice(image_paths)
        img_file = Path(img_path)
        
        if not img_file.exists():
            logger.warning(f"图片文件不存在: {img_path}")
            return
        
        try:
            # 添加图片段落
            paragraph = doc.add_paragraph()
            paragraph.alignment = 1  # 居中对齐
            run = paragraph.add_run()
            
            # 获取图片原始尺寸
            with Image.open(str(img_file)) as img:
                img_width, img_height = img.size
                aspect_ratio = img_height / img_width
            
            # Word A4 文档可用宽度约为 16cm（左右边距各2.54cm，总宽21cm）
            # 设置图片宽度为可用宽度的 90%，即 14.4cm
            max_width = Cm(14.4)
            
            # 插入图片，自动按比例调整高度
            picture = run.add_picture(str(img_file), width=max_width)
            
            # 提取文件名（去掉后缀）作为 Alt Text
            alt_text = img_file.stem  # 自动去掉扩展名
            
            # 设置图片的 Alt Text（替代文本）
            # 这是 SEO 的关键部分
            # 修正：使用 ._inline 而不是 .inline
            inline = picture._inline 
            docPr = inline.docPr
            docPr.set('descr', alt_text)  # 设置描述（Alt Text）
            docPr.set('title', alt_text)  # 同时设置标题
            
            logger.info(f"列 {col_idx+1} 插入图片: {img_file.name}, Alt Text: {alt_text}, 宽度: 14.4cm")
            
        except Exception as e:
            logger.error(f"插入图片失败: {img_path}, 错误: {e}")
    
    def _check_and_insert_comparison_table(self, doc, col_idx: int, current_content: str, row_data: list):
        """检查并插入对比表图片（支持多任务）
        
        Args:
            doc: Document对象
            col_idx: 列索引
            current_content: 当前列的内容
            row_data: 整行数据（用于提取品牌）
        """
        try:
            # 导入对比表模块
            from ..core.comparison_image_generator import ComparisonTableImageGenerator
            from ..database.comparison_db_manager import ComparisonDBManager
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import os
            
            # 初始化管理器
            comparison_db = ComparisonDBManager()
            comparison_generator = ComparisonTableImageGenerator()
            
            # 加载全局配置
            global_config = comparison_db.get_config('insert_strategy')
            if not global_config:
                logger.debug("未找到全局配置")
                return
            
            # 获取所有类目
            categories = comparison_db.get_all_categories()
            if not categories:
                logger.warning("未找到对比表类目")
                return
            
            # 使用第一个类目
            category = categories[0]
            
            # 获取该类目下的所有任务（按排序）
            tasks = comparison_db.get_tasks_by_category(category.id)
            if not tasks:
                logger.debug("该类目下没有任务")
                return
            
            # 提取文章中的品牌（所有任务共用）
            full_text = " ".join([str(c) for c in row_data if c])
            mentioned_brands = self._extract_mentioned_brands(comparison_db, full_text)
            
            # 遍历所有任务，检查是否需要插入
            for task in tasks:
                should_insert = False
                insert_reason = ""
                
                # 判断是否需要插入
                if task.insert_mode == 'column':
                    # 按列插入
                    if col_idx == task.insert_column - 1:
                        should_insert = True
                        insert_reason = f"任务'{task.task_name}': 按列插入（列{col_idx}）"
                
                elif task.insert_mode == 'anchor':
                    # 智能锚点
                    if task.insert_anchor_text and task.insert_anchor_text in current_content:
                        should_insert = True
                        insert_reason = f"任务'{task.task_name}': 锚点匹配（'{task.insert_anchor_text}'）"
                
                if not should_insert:
                    continue
                
                logger.info(f"✓ 触发对比表插入: {insert_reason}")
                
                # 获取任务的参数选择
                selected_param_ids = comparison_db.get_task_parameters(task.id)
                if not selected_param_ids:
                    logger.warning(f"任务'{task.task_name}'未选择任何参数，跳过")
                    continue
                
                # 获取任务的样式配置
                style_config = task.get_style_dict()
                if not style_config:
                    # 使用默认样式
                    style_config = {
                        'header_bg_color': '#4472C4',
                        'header_text_color': '#FFFFFF',
                        'own_brand_bg_color': '#FFF2CC',
                        'border_width': 1.5,
                        'image_width': 15,
                        'dpi': 300,
                        'font_name': 'Microsoft YaHei',
                        'font_size': 10
                    }
                
                # 生成图片
                image_path = comparison_generator.generate_from_category(
                    db_manager=comparison_db,
                    category_id=category.id,
                    mentioned_brands=mentioned_brands,
                    style_config=style_config,
                    insert_config=global_config,
                    selected_parameter_ids=selected_param_ids
                )
                
                # 插入图片
                if image_path and os.path.exists(image_path):
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run()
                    
                    image_width = style_config.get('image_width', 15)
                    run.add_picture(image_path, width=Inches(image_width / 2.54))
                    
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    logger.info(f"✓ 对比表图片已插入: {task.task_name}")
                else:
                    logger.warning(f"对比表图片生成失败: {task.task_name}")
        
        except ImportError as e:
            logger.debug(f"对比表功能不可用: {e}")
        except Exception as e:
            logger.error(f"插入对比表失败: {e}")
            import traceback
            logger.error(traceback.format_exc())
    
    def _extract_mentioned_brands(self, comparison_db, text: str) -> list:
        """提取文章中提及的品牌（仅完整匹配）
        
        Args:
            comparison_db: 数据库管理器
            text: 文档文本
            
        Returns:
            品牌名称列表
        """
        mentioned_brands = []
        
        categories = comparison_db.get_all_categories()
        for category in categories:
            brands = comparison_db.get_brands_by_category(category.id)
            for brand in brands:
                brand_name = brand.name
                
                # 仅完整匹配（精确匹配完整品牌名）
                if brand_name in text:
                    mentioned_brands.append(brand_name)
                    logger.debug(f"品牌完整匹配: {brand_name}")
        
        logger.info(f"识别到的品牌: {mentioned_brands if mentioned_brands else '无'}")
        return mentioned_brands
    
    def _apply_heading_style(self, paragraph, level: int):
        """应用标题样式
        
        Args:
            paragraph: 段落对象
            level: 标题级别 (1-4)
        """
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
        
        # 字号映射（中国公文标准）
        font_sizes = {
            1: 24,  # 小一号
            2: 18,  # 小二号
            3: 16,  # 小三号
            4: 14   # 四号
        }
        
        font_size = font_sizes.get(level, 16)
        
        # 设置段落格式
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = 1.5  # 1.5倍行距
        paragraph_format.space_after = Pt(10)  # 段后10pt
        
        # 如果段落为空，添加一个run
        if not paragraph.runs:
            paragraph.add_run()
        
        # 对每个run应用样式
        for run in paragraph.runs:
            # 设置字体
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(font_size)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
            
            # 强制设置中文字体（核心修复）
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    def _apply_body_style(self, paragraph):
        """应用正文样式"""
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
        
        # 设置段落格式
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = 1.5  # 1.5倍行距
        paragraph_format.space_after = Pt(10)  # 段后10pt
        
        # 如果段落为空，添加一个run
        if not paragraph.runs:
            paragraph.add_run()
        
        # 对每个run应用样式
        for run in paragraph.runs:
            # 设置字体
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(12)  # 小四号
            run.font.bold = False
            run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
            
            # 强制设置中文字体（核心修复）
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    def _apply_column_shuffling_strategies(self, row_data: list) -> list:
        """
        应用混排策略（只保留/删除指定列，不改变内容）
        
        Args:
            row_data: 行数据
            
        Returns:
            应用策略后的行数据
        """
        import random
        import copy
        
        result_row = copy.deepcopy(row_data)
        
        # 应用每个策略
        for strategy in self.config.shuffling_strategies:
            # 将列索引转换为0-based（策略中存储的是1-based，即用户看到的"第1列"、"第2列"）
            columns = [col - 1 for col in strategy.columns if col > 0]
            
            logger.info(f"应用策略 '{strategy.name}': 原始列号 {strategy.columns} -> 0-based索引 {columns}, 分组大小={strategy.group_size}, 保留组数={strategy.keep_count}")
            
            # 验证列索引范围
            valid_columns = [col for col in columns if 0 <= col < len(result_row)]
            if len(valid_columns) != len(columns):
                logger.warning(f"策略 '{strategy.name}' 部分列索引超出范围，过滤后: {valid_columns}")
            
            if not valid_columns:
                logger.warning(f"策略 '{strategy.name}' 没有有效的列索引，跳过")
                continue
            
            # 分组
            groups = []
            for i in range(0, len(valid_columns), strategy.group_size):
                group = valid_columns[i:i + strategy.group_size]
                # 只保留完整的组
                if len(group) == strategy.group_size:
                    groups.append(group)
                else:
                    logger.debug(f"跳过不完整的组: {group}")
            
            if not groups:
                logger.warning(f"策略 '{strategy.name}' 无法形成完整分组，跳过")
                continue
            
            logger.debug(f"策略 '{strategy.name}' 共分为 {len(groups)} 组: {groups}")
            
            # 随机选择保留的组
            keep_count = min(strategy.keep_count, len(groups))
            kept_groups = random.sample(groups, keep_count)
            
            logger.debug(f"随机保留 {keep_count} 组: {kept_groups}")
            
            # 如果需要打乱顺序
            if strategy.shuffle_order:
                random.shuffle(kept_groups)
                logger.debug(f"打乱顺序后: {kept_groups}")
            
            # 展开为列索引集合
            kept_columns = set()
            for group in kept_groups:
                kept_columns.update(group)
            
            # 删除未保留的列（设为空）
            deleted_columns = []
            for col in valid_columns:
                if col not in kept_columns:
                    result_row[col] = ""
                    deleted_columns.append(col)
            
            logger.info(f"策略 '{strategy.name}': 保留列 {sorted(kept_columns)}, 删除列 {sorted(deleted_columns)}")
        
        return result_row
    
    def _apply_shuffling_strategies(self, grid_data: list, base_row_idx: int) -> list:
        """
        应用混排策略，生成新的行数据（旧方法，保留用于按行生成模式）
        """
        import random
        import copy
        
        # 获取基础行
        base_row = list(grid_data[base_row_idx])
        result_row = copy.deepcopy(base_row)
        
        # 如果没有配置策略，直接返回基础行
        if not self.config.shuffling_strategies:
            return result_row
        
        # 应用每个策略
        for strategy in self.config.shuffling_strategies:
            result_row = self._apply_single_strategy(grid_data, result_row, strategy)
        
        return result_row
    
    def _apply_single_strategy(self, grid_data: list, current_row: list, strategy) -> list:
        """应用单个混排策略
        
        逻辑说明：
        1. 将指定的列按group_size分组
        2. 随机选择keep_count个组保留
        3. 未被选中的组对应的列删除（设为空）
        4. 如果shuffle_order=True，打乱保留的组之间的位置顺序
        """
        import random
        import copy
        
        result_row = copy.deepcopy(current_row)
        
        # 将列索引转换为0-based（策略中是1-based）
        columns = [col - 1 for col in strategy.columns if col > 0]
        
        # 验证列索引是否在范围内
        valid_columns = [col for col in columns if col < len(current_row)]
        if not valid_columns:
            logger.warning(f"策略 '{strategy.name}' 的列索引超出范围，跳过")
            return result_row
        
        # 分组
        groups = []
        for i in range(0, len(valid_columns), strategy.group_size):
            group = valid_columns[i:i + strategy.group_size]
            if len(group) == strategy.group_size:  # 只保留完整的组
                groups.append(group)
        
        if not groups:
            logger.warning(f"策略 '{strategy.name}' 无法形成完整分组，跳过")
            return result_row
        
        # 随机选择指定数量的组（保留的组）
        keep_count = min(strategy.keep_count, len(groups))
        selected_groups = random.sample(groups, keep_count)
        
        # 记录原始组的位置和内容
        original_positions = []  # [(组位置, 组列索引)]
        for group in selected_groups:
            original_positions.append(group)
        
        # 如果需要打乱保留组之间的顺序
        if strategy.shuffle_order:
            # 保存选中组的内容
            group_contents = []
            for group in selected_groups:
                content = [result_row[col_idx] for col_idx in group]
                group_contents.append(content)
            
            # 打乱内容顺序
            random.shuffle(group_contents)
            
            # 将打乱后的内容重新分配到原来的组位置
            for group_idx, group in enumerate(selected_groups):
                for col_idx_in_group, col_idx in enumerate(group):
                    result_row[col_idx] = group_contents[group_idx][col_idx_in_group]
        
        # 收集所有涉及策略的列索引
        all_strategy_columns = set(valid_columns)
        
        # 收集要保留的列索引
        keep_columns = set()
        for group in selected_groups:
            keep_columns.update(group)
        
        # 找出要删除的列（在策略范围内但未被选中的列）
        columns_to_remove = all_strategy_columns - keep_columns
        
        # 将要删除的列设置为空字符串
        for col_idx in columns_to_remove:
            result_row[col_idx] = ""
        
        logger.debug(f"应用策略 '{strategy.name}': "
                    f"总共 {len(groups)} 组, 保留 {keep_count} 组, "
                    f"删除 {len(columns_to_remove)} 列, "
                    f"打乱顺序: {strategy.shuffle_order}")
        
        return result_row
    
    def _apply_bold_keywords(self, paragraph, keywords: list):
        """应用加粗关键词
        
        Args:
            paragraph: 段落对象
            keywords: 关键词列表
        """
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
        
        if not keywords or not paragraph.text:
            return
        
        # 获取原始文本
        original_text = paragraph.text
        
        # 清空段落的所有runs
        for run in paragraph.runs:
            run.text = ''
        
        # 重新构建段落，对关键词加粗
        current_pos = 0
        text_length = len(original_text)
        
        while current_pos < text_length:
            # 查找最近的关键词
            nearest_keyword = None
            nearest_pos = text_length
            
            for keyword in keywords:
                pos = original_text.find(keyword, current_pos)
                if pos != -1 and pos < nearest_pos:
                    nearest_pos = pos
                    nearest_keyword = keyword
            
            if nearest_keyword:
                # 添加关键词之前的普通文本
                if nearest_pos > current_pos:
                    run = paragraph.add_run(original_text[current_pos:nearest_pos])
                    run.font.name = 'Microsoft YaHei'
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                
                # 添加加粗的关键词
                run = paragraph.add_run(nearest_keyword)
                run.font.name = 'Microsoft YaHei'
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                
                current_pos = nearest_pos + len(nearest_keyword)
            else:
                # 添加剩余的普通文本
                run = paragraph.add_run(original_text[current_pos:])
                run.font.name = 'Microsoft YaHei'
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                break
    
    def _on_generate_complete(self, save_dir: str):
        """生成完成（废弃，已整合到 _on_generate 中）"""
        pass
    
    def _on_ai_title_dialog(self):
        """打开 AI 标题生成对话框"""
        dialog = AITitleDialog(self.config, self)
        if dialog.exec():
            # 用户点击了"确认并使用"
            titles = dialog.get_titles()
            title_format = dialog.get_title_format()
            
            if not titles:
                return
            
            # 保存到标题队列
            self.ai_title_queue = titles
            self.ai_title_format = title_format
            
            # 强制设置生成数量为标题数量
            self.toolbar.count_spin.setValue(len(titles))
            
            # 🔒 锁定生成数量输入框（AI标题模式）
            self.toolbar.count_spin.setEnabled(False)
            self.toolbar.count_spin.setToolTip("AI标题模式下，生成数量已自动锁定")
            
            # 强制切换到"随机混排"模式
            self.toolbar.mode_combo.setCurrentIndex(1)
            
            from qfluentwidgets import InfoBar, InfoBarPosition
            InfoBar.success(
                title='✅ AI标题已就绪',
                content=f'已加载 {len(titles)} 个标题，将自动生成 {len(titles)} 篇文章',
                orient=Qt.Orientation.Horizontal,
                isClosable=False,
                position=InfoBarPosition.BOTTOM_RIGHT,
                duration=3000,
                parent=self
            )
            
            logger.info(f"AI 标题队列已设置: {len(titles)} 个标题，格式: {title_format}，生成数量已锁定")
    
    def _on_ai_rewrite_dialog(self):
        """打开 AI 内容改写对话框（支持多列）"""
        # 获取当前网格数据
        grid_data = self.smart_grid.get_grid_data()
        
        if not grid_data:
            from qfluentwidgets import InfoBar, InfoBarPosition
            InfoBar.warning(
                title='无数据',
                content='工作区没有数据，请先导入或添加内容',
                orient=Qt.Orientation.Horizontal,
                isClosable=True,
                position=InfoBarPosition.BOTTOM_RIGHT,
                duration=3000,
                parent=self
            )
            return
        
        # 打开对话框
        dialog = AIRewriteDialog(self.config, grid_data, self)
        if dialog.exec():
            # 用户点击了"确认追加"
            rewritten_results = dialog.get_rewritten_results()
            
            if not rewritten_results:
                return
            
            # 批量追加内容到多个列
            total_count = 0
            for column_index, contents in rewritten_results.items():
                if contents:
                    self._append_contents_to_column(column_index, contents)
                    total_count += len(contents)
            
            from qfluentwidgets import InfoBar, InfoBarPosition
            InfoBar.success(
                title='追加成功',
                content=f'已向 {len(rewritten_results)} 列追加共 {total_count} 个新内容',
                orient=Qt.Orientation.Horizontal,
                isClosable=False,
                position=InfoBarPosition.BOTTOM_RIGHT,
                duration=3000,
                parent=self
            )
            
            logger.info(f"AI 改写内容已追加: {len(rewritten_results)} 列，共 {total_count} 个新内容")
    
    def _find_last_row_in_column(self, column_index: int) -> int:
        """
        找到指定列的最后一个非空单元格的行号
        
        Args:
            column_index: 列索引
            
        Returns:
            最后一个非空单元格的行号，如果列为空则返回 -1
        """
        table = self.smart_grid.table
        last_row = -1
        
        # 从上往下扫描该列
        for row in range(table.rowCount()):
            item = table.item(row, column_index)
            if item and item.text().strip():
                last_row = row
        
        return last_row
    
    def _append_contents_to_column(self, column_index: int, contents: list):
        """
        将内容追加到指定列的底部（紧接该列最后一行）
        
        Args:
            column_index: 列索引
            contents: 要追加的内容列表
        """
        from PyQt6.QtWidgets import QTableWidgetItem
        
        table = self.smart_grid.table
        
        # 找到该列最后一个非空单元格的行号
        last_row_in_column = self._find_last_row_in_column(column_index)
        
        # 计算起始行（该列最后一行的下一行）
        start_row = last_row_in_column + 1
        
        # 计算需要的总行数
        needed_rows = start_row + len(contents)
        
        # 如果需要的行数超过当前表格行数，扩展表格
        if needed_rows > table.rowCount():
            table.setRowCount(needed_rows)
        
        # 追加内容（只填充指定列，其他列保持空白）
        for i, content in enumerate(contents):
            row_index = start_row + i
            item = QTableWidgetItem(content)
            table.setItem(row_index, column_index, item)
        
        logger.debug(f"已向列 {column_index + 1} 追加 {len(contents)} 行内容（从第 {start_row + 1} 行开始）")
    
    def _on_strategy_config(self):
        """打开混排策略配置对话框"""
        dialog = StrategyConfigDialog(self.config, self._validate_strategy_columns, self)
        if dialog.exec():
            # 用户点击了"保存生效"
            # 配置已在对话框中保存，这里只需要刷新界面
            self.strategy_panel._update_strategy_count()
            
            from qfluentwidgets import InfoBar, InfoBarPosition
            InfoBar.success(
                title='配置已更新',
                content=f'混排策略配置已生效',
                orient=Qt.Orientation.Horizontal,
                isClosable=False,
                position=InfoBarPosition.BOTTOM_RIGHT,
                duration=2000,
                parent=self
            )
            
            logger.info("混排策略配置已更新")
    
    def _on_numbering_group_config(self):
        """打开序号分组配置对话框"""
        from .dialogs.numbering_group_dialog import NumberingGroupDialog
        from qfluentwidgets import InfoBar, InfoBarPosition
        
        dialog = NumberingGroupDialog(self.config, self)
        if dialog.exec():
            # 配置已保存
            logger.info(f"序号分组配置已更新: {self.config.numbering_groups}")
            InfoBar.success(
                title='配置已更新',
                content=f'已保存 {len(self.config.numbering_groups)} 个序号分组',
                orient=Qt.Orientation.Horizontal,
                isClosable=False,
                position=InfoBarPosition.BOTTOM_RIGHT,
                duration=2000,
                parent=self
            )
    
    def _on_seo_config(self):
        """打开 SEO 核心词配置对话框"""
        dialog = SEOSettingDialog(self.config, self)
        if dialog.exec():
            # 用户点击了"保存并生效"
            # 配置已在对话框中保存，这里只需要刷新界面
            self.strategy_panel.update_seo_status()
            
            from qfluentwidgets import InfoBar, InfoBarPosition
            InfoBar.success(
                title='配置已更新',
                content=f'SEO 核心词配置已生效',
                orient=Qt.Orientation.Horizontal,
                isClosable=False,
                position=InfoBarPosition.BOTTOM_RIGHT,
                duration=2000,
                parent=self
            )
            
            logger.info(f"SEO 核心词配置已更新: {self.config.target_keywords}")
    
    def _on_dedup_config(self):
        """打开历史查重配置对话框"""
        from .dialogs.dedup_config_dialog import DedupConfigDialog
        
        dialog = DedupConfigDialog(self.config, self)
        if dialog.exec():
            # 用户点击了"保存"，配置已在对话框中保存
            from qfluentwidgets import InfoBar, InfoBarPosition
            InfoBar.success(
                title='配置已更新',
                content=f'历史查重配置已生效',
                orient=Qt.Orientation.Horizontal,
                isClosable=False,
                position=InfoBarPosition.BOTTOM_RIGHT,
                duration=2000,
                parent=self
            )
            
            logger.info(f"历史查重配置已更新: 启用={self.config.dedup_enabled}, 阈值={self.config.dedup_similarity_threshold}")
    
    def _on_ai_title(self, keyword: str, prompt: str):
        """AI 生成标题"""
        logger.info(f"AI 生成标题: keyword={keyword}, prompt={prompt}")
        from qfluentwidgets import InfoBar, InfoBarPosition
        from PyQt6.QtWidgets import QApplication, QTableWidgetItem
        from ..ai.api_client import AIClient
        
        # 检查 API 配置
        if not self.config.api_config.api_key:
            InfoBar.error(
                title='配置错误',
                content='请先在设置中配置 API Key',
                orient=Qt.Orientation.Horizontal,
                isClosable=True,
                position=InfoBarPosition.BOTTOM_RIGHT,
                duration=3000,
                parent=self
            )
            return
        
        # 显示加载状态
        QApplication.setOverrideCursor(Qt.CursorShape.WaitCursor)
        
        row_count = self.smart_grid.get_active_row_count()
        if row_count == 0:
            row_count = max(1, self.smart_grid.table.rowCount())
        
        InfoBar.info(
            title='AI 助手',
            content=f"正在为 '{keyword}' 生成 {row_count} 个标题...",
            orient=Qt.Orientation.Horizontal,
            isClosable=False,
            position=InfoBarPosition.BOTTOM_RIGHT,
            duration=2000,
            parent=self
        )
        
        try:
            # 初始化 AI 客户端
            ai_client = AIClient(self.config.api_config)
            
            # 调用 AI 生成标题
            titles = ai_client.generate_titles(
                keyword=keyword,
                count=row_count,
                custom_prompt=prompt if prompt else ""
            )
            
            if not titles:
                QApplication.restoreOverrideCursor()
                InfoBar.error(
                    title='生成失败',
                    content='AI 未返回任何标题，请检查网络或 API 配置',
                    orient=Qt.Orientation.Horizontal,
                    isClosable=True,
                    position=InfoBarPosition.BOTTOM_RIGHT,
                    duration=3000,
                    parent=self
                )
                return
            
            # 确保表格有足够的行和列
            if self.smart_grid.table.columnCount() == 0:
                self.smart_grid.table.setColumnCount(1)
                self.smart_grid._setup_headers()
            if self.smart_grid.table.rowCount() < len(titles):
                self.smart_grid.table.setRowCount(len(titles))
            
            # 填充标题到第一列
            for i, title in enumerate(titles):
                self.smart_grid.table.setItem(i, 0, QTableWidgetItem(title))
            
            self.smart_grid.empty_hint.hide()
            self.smart_grid.table.show()
            self.smart_grid._update_column_control_positions()
            
            # 恢复光标
            QApplication.restoreOverrideCursor()
            
            InfoBar.success(
                title='生成完成',
                content=f'已生成 {len(titles)} 个标题并导入到第一列',
                orient=Qt.Orientation.Horizontal,
                isClosable=False,
                position=InfoBarPosition.BOTTOM_RIGHT,
                duration=3000,
                parent=self
            )
            
            logger.info(f"AI 标题已导入工作区: {len(titles)} 个")
            
        except Exception as e:
            QApplication.restoreOverrideCursor()
            logger.error(f"AI 生成标题失败: {e}")
            InfoBar.error(
                title='生成失败',
                content=f'错误: {str(e)}',
                orient=Qt.Orientation.Horizontal,
                isClosable=True,
                position=InfoBarPosition.BOTTOM_RIGHT,
                duration=5000,
                parent=self
            )
    
    def _on_material_selected(self, material_id: int):
        """素材被选中"""
        logger.debug(f"素材被选中: ID={material_id}")
    
    def _on_grid_data_changed(self):
        """网格数据变化"""
        logger.debug("网格数据已变化")
        # 更新按钮状态
        self._update_button_states()
    
    def _update_button_states(self):
        """
        更新右侧功能面板按钮状态
        根据工作区是否有数据来启用/禁用相关按钮
        """
        # 检查工作区是否有数据
        grid_data = self.smart_grid.get_grid_data()
        has_data = len(grid_data) > 0
        
        # 更新策略面板的按钮状态
        self.strategy_panel.update_button_states(has_data)
        
        logger.debug(f"按钮状态已更新: has_data={has_data}, 数据行数={len(grid_data) if grid_data else 0}")
