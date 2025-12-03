"""
Word 文档生成引擎
核心功能：将网格数据转换为格式化的 Word 文档
"""

import os
from pathlib import Path
from typing import List, Dict, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from loguru import logger

from .spintax_parser import SpintaxParser
from .image_processor import ImageProcessor
from .shuffle_engine import ShuffleEngine, SmartShuffle
from ..config.settings import ProfileConfig
from ..utils.file_handler import FileHandler
from ..database.db_manager import DatabaseManager
from ..database.fingerprint_manager import FingerprintManager
from .simhash_deduplicator import ContentDeduplicator

# 对比表功能（可选依赖）
try:
    from .comparison_image_generator import ComparisonTableImageGenerator
    from ..database.comparison_db_manager import ComparisonDBManager
    COMPARISON_TABLE_AVAILABLE = True
except ImportError as e:
    logger.warning(f"对比表功能不可用（缺少依赖）: {e}")
    ComparisonTableImageGenerator = None
    ComparisonDBManager = None
    COMPARISON_TABLE_AVAILABLE = False


class DocumentGenerator:
    """Word 文档生成器"""
    
    def __init__(self, config: ProfileConfig):
        """
        初始化文档生成器
        
        Args:
            config: 配置对象
        """
        self.config = config
        self.spintax_parser = SpintaxParser()
        self.image_processor = ImageProcessor()
        
        # 初始化混排引擎
        if config.shuffling_strategies:
            self.shuffle_engine = ShuffleEngine(config.shuffling_strategies)
        else:
            self.shuffle_engine = None
        
        # 初始化查重器
        if config.dedup_enabled:
            logger.info("✓ 历史查重功能已启用")
            db_manager = DatabaseManager()
            fp_manager = FingerprintManager(db_manager)
            
            dedup_config = {
                'enabled': config.dedup_enabled,
                'max_distance': config.get_dedup_max_distance(),
                'max_retries': config.dedup_max_retries,
                'retention_days': config.dedup_retention_days,
            }
            
            self.deduplicator = ContentDeduplicator(fp_manager, dedup_config)
            logger.info(f"  - 相似度阈值: {config.dedup_similarity_threshold*100:.0f}% (海明距离≤{config.get_dedup_max_distance()})")
            logger.info(f"  - 最大重试: {config.dedup_max_retries} 次")
            logger.info(f"  - 项目范围: {'全局' if config.dedup_cross_project else config.dedup_current_project}")
        else:
            self.deduplicator = None
            logger.debug("历史查重功能未启用")
        
        # 初始化对比表生成器和数据库管理器（如果可用）
        if COMPARISON_TABLE_AVAILABLE:
            logger.info("✓ 对比表功能可用，正在初始化...")
            self.comparison_generator = ComparisonTableImageGenerator()
            self.comparison_db = ComparisonDBManager()
            self.comparison_table_config = self._load_comparison_config()
            if self.comparison_table_config:
                logger.info("✓ 对比表配置已加载")
            else:
                logger.warning("⚠ 对比表配置未设置（请在数据库界面配置）")
        else:
            logger.warning("⚠ 对比表功能不可用（matplotlib未安装）")
            self.comparison_generator = None
            self.comparison_db = None
            self.comparison_table_config = None
    
    def generate_by_row(
        self,
        grid_data: List[List[str]],
        output_dir: str = "output"
    ) -> List[str]:
        """
        按行生成模式：每行数据生成一个文档
        
        Args:
            grid_data: 网格数据（二维列表）
            output_dir: 输出目录
            
        Returns:
            生成的文件路径列表
        """
        FileHandler.ensure_directory(output_dir)
        generated_files = []
        
        for row_idx, row_data in enumerate(grid_data):
            try:
                # 创建文档
                doc = self._create_document()
                
                # 添加内容
                self._add_row_content(doc, row_data, row_idx)
                
                # 生成文件名
                filename = self._generate_filename(row_data, row_idx)
                filepath = os.path.join(output_dir, filename)
                
                # 保存文档
                if FileHandler.save_word(doc, filepath):
                    generated_files.append(filepath)
                    logger.info(f"生成文档 {row_idx + 1}/{len(grid_data)}: {filename}")
                
            except Exception as e:
                logger.error(f"生成第 {row_idx + 1} 行文档失败: {e}")
        
        logger.info(f"按行生成完成，共 {len(generated_files)} 个文档")
        return generated_files
    
    def generate_by_shuffle(
        self,
        grid_data: List[List[str]],
        count: int,
        output_dir: str = "output",
        columns_data: Optional[List[List[str]]] = None
    ) -> List[str]:
        """
        随机混排模式：随机组合生成指定数量的文档
        
        Args:
            grid_data: 网格数据（二维列表，按行组织）
            count: 生成数量
            output_dir: 输出目录
            columns_data: 可选，直接传入按列组织的数据（优先使用）
            
        Returns:
            生成的文件路径列表
        """
        FileHandler.ensure_directory(output_dir)
        generated_files = []
        
        # 查重统计
        duplicate_count = 0
        retry_count = 0
        
        # 获取项目名称（用于查重）
        project_name = self.config.get_dedup_project_name() if self.deduplicator else None
        
        # 如果直接提供了列数据，使用它；否则转置行数据
        if columns_data is not None:
            logger.info("使用直接提供的列数据（避免转置）")
        else:
            logger.info("转置行数据为列数据")
            columns_data = self._transpose_grid(grid_data)
        
        total_columns = len(columns_data)
        
        # 为每列创建智能轮播器（只包含非空内容）
        column_shufflers = []
        for col_data in columns_data:
            # 过滤空值，只保留有效内容
            valid_items = [item for item in col_data if item and item.strip()]
            if valid_items:
                shuffler = SmartShuffle(len(valid_items))
                column_shufflers.append((valid_items, shuffler))
            else:
                column_shufflers.append(([], None))
        
        logger.info(f"初始化混排器：共 {total_columns} 列")
        for i, (items, shuffler) in enumerate(column_shufflers):
            logger.debug(f"列 {i+1}: {len(items)} 个有效内容")
        
        doc_idx = 0
        attempts = 0
        max_total_attempts = count * (self.config.dedup_max_retries if self.deduplicator else 1)
        
        while doc_idx < count and attempts < max_total_attempts:
            attempts += 1
            
            try:
                # 使用智能轮播器随机选择每列的一行
                selected_row = []
                for valid_items, shuffler in column_shufflers:
                    if shuffler and valid_items:
                        # 使用轮播器获取下一个索引（避免重复，用完后自动重置）
                        index = shuffler.get_next_index()
                        selected_row.append(valid_items[index])
                    else:
                        # 该列没有有效内容，使用空字符串
                        selected_row.append("")
                
                # 应用混排策略
                if self.shuffle_engine:
                    keep_map = self.shuffle_engine.execute(total_columns)
                    selected_row = [
                        cell if keep_map.get(i, True) else ""
                        for i, cell in enumerate(selected_row)
                    ]
                
                # 提取全文本（用于查重）
                full_text = self._extract_full_text(selected_row)
                
                # 查重检查
                if self.deduplicator:
                    is_duplicate, dup_info = self.deduplicator.check_duplicate(
                        text=full_text,
                        source_project=project_name
                    )
                    
                    if is_duplicate:
                        duplicate_count += 1
                        retry_count += 1
                        
                        similarity = dup_info.get('similarity_percent', 100)
                        logger.warning(
                            f"⚠ 检测到重复内容 (相似度: {similarity:.1f}%), "
                            f"重试 {retry_count}/{self.config.dedup_max_retries}"
                        )
                        
                        # 检查是否超过最大重试次数
                        if retry_count >= self.config.dedup_max_retries:
                            logger.error(
                                f"✗ 文档 {doc_idx + 1} 超过最大重试次数，跳过"
                            )
                            doc_idx += 1  # 跳过这篇
                            retry_count = 0
                        
                        continue  # 重新生成
                
                # 通过查重，开始生成文档
                retry_count = 0  # 重置重试计数
                
                # 创建文档
                doc = self._create_document()
                
                # 添加内容
                self._add_row_content(doc, selected_row, doc_idx)
                
                # 生成文件名
                filename = self._generate_filename(selected_row, doc_idx)
                filepath = os.path.join(output_dir, filename)
                
                # 保存文档
                if FileHandler.save_word(doc, filepath):
                    generated_files.append(filepath)
                    
                    # 保存成功后，将指纹写入数据库
                    if self.deduplicator:
                        relative_path = os.path.relpath(filepath, output_dir)
                        self.deduplicator.add_content_fingerprint(
                            text=full_text,
                            source_project=project_name or "default",
                            document_path=relative_path
                        )
                        logger.debug(f"✓ 指纹已记录: {filename}")
                    
                    logger.info(f"✓ 生成文档 {doc_idx + 1}/{count}: {filename}")
                    doc_idx += 1
                
            except Exception as e:
                logger.error(f"生成第 {doc_idx + 1} 个文档失败: {e}")
                doc_idx += 1  # 继续下一篇
                retry_count = 0
        
        # 生成完成，输出统计
        logger.info(f"混排生成完成，共 {len(generated_files)} 个文档")
        
        if self.deduplicator and duplicate_count > 0:
            logger.warning(f"查重统计: 拦截 {duplicate_count} 次重复, 总尝试 {attempts} 次")
        
        return generated_files
    
    def _create_document(self) -> Document:
        """
        创建 Word 文档（基于模板或空白）
        
        Returns:
            Document 对象
        """
        template_path = self.config.template_path
        return FileHandler.create_word_from_template(template_path)
    
    def _add_row_content(self, doc: Document, row_data: List[str], row_idx: int):
        """
        将一行数据添加到文档
        
        Args:
            doc: Document 对象
            row_data: 行数据
            row_idx: 行索引
        """
        # 收集文档全文，用于品牌识别
        full_text = " ".join(row_data)
        
        for col_idx, cell_content in enumerate(row_data):
            if not cell_content or not cell_content.strip():
                continue
            
            # 获取列类型
            col_type = self.config.get_column_type(col_idx)
            
            # 忽略此列
            if col_type == 'Ignore':
                continue
            
            # 解析 Spintax
            parsed_content = self.spintax_parser.parse(cell_content)
            
            # 根据类型添加内容
            if col_type == 'H1':
                self._add_heading(doc, parsed_content, level=1)
            elif col_type == 'H2':
                self._add_heading(doc, parsed_content, level=2)
            elif col_type == 'H3':
                self._add_heading(doc, parsed_content, level=3)
            elif col_type == 'List':
                self._add_list_item(doc, parsed_content)
            else:  # Body
                self._add_paragraph(doc, parsed_content)
            
            # 插入图片（如果该列绑定了图片文件夹）
            image_path = self.config.get_image_path(col_idx)
            if image_path:
                self._add_image(doc, image_path, row_idx)
            
            # 检查是否需要插入对比表图片
            self._check_and_insert_comparison_table(doc, col_idx, parsed_content, full_text)
    
    def _check_and_insert_comparison_table(
        self, 
        doc: Document, 
        col_idx: int, 
        current_content: str,
        full_text: str
    ):
        """
        检查是否需要插入对比表图片
        
        Args:
            doc: Document 对象
            col_idx: 当前列索引
            current_content: 当前段落内容
            full_text: 文档全文
        """
        # 调试日志：方法被调用
        logger.debug(f"_check_and_insert_comparison_table 被调用: col_idx={col_idx}")
        
        # 如果对比表功能不可用，直接返回
        if not COMPARISON_TABLE_AVAILABLE:
            logger.warning("对比表功能不可用: COMPARISON_TABLE_AVAILABLE=False (可能缺少matplotlib)")
            return
        
        if not self.comparison_table_config:
            logger.warning("对比表配置未加载: comparison_table_config=None (请在数据库界面配置)")
            return
        
        insert_config = self.comparison_table_config.get('insert_strategy')
        style_config = self.comparison_table_config.get('table_style')
        
        if not insert_config:
            logger.warning("未找到插入策略配置 (请在数据库界面点击'插入策略配置')")
            return
        
        logger.debug(f"插入配置: 模式={insert_config.get('insert_mode')}, 列号={insert_config.get('insert_column')}, 锚点={insert_config.get('insert_anchor_text')}")
        
        should_insert = False
        insert_reason = ""
        
        # 判断是否需要插入
        if insert_config['insert_mode'] == 'column':
            # 按列插入：在第N列后插入
            target_column = insert_config['insert_column'] - 1
            logger.debug(f"按列模式: 目标列={target_column}, 当前列={col_idx}")
            if col_idx == target_column:
                should_insert = True
                insert_reason = f"按列插入（列索引={col_idx}）"
        
        elif insert_config['insert_mode'] == 'anchor':
            # 智能锚点：检测到特定文本后插入
            anchor_text = insert_config['insert_anchor_text']
            logger.debug(f"锚点模式: 锚点文本='{anchor_text}', 当前内容前50字='{current_content[:50]}'")
            if anchor_text and anchor_text in current_content:
                should_insert = True
                insert_reason = f"智能锚点匹配（锚点文本='{anchor_text}'）"
        
        if not should_insert:
            logger.debug(f"列{col_idx}: 不满足插入条件（模式={insert_config['insert_mode']}）")
            return
        
        logger.info(f"✓ 触发对比表插入: {insert_reason}")
        
        # 提取文档中提及的品牌
        mentioned_brands = self._extract_mentioned_brands(full_text)
        
        # 获取所有类目
        categories = self.comparison_db.get_all_categories()
        if not categories:
            logger.warning("未找到对比表类目，无法生成对比表")
            return
        
        # 使用第一个类目（后续可以扩展为智能匹配类目）
        category = categories[0]
        logger.info(f"使用类目: {category.name} (ID={category.id})")
        
        try:
            # 生成对比表图片
            logger.info(f"开始生成对比表图片...")
            image_path = self.comparison_generator.generate_from_category(
                db_manager=self.comparison_db,
                category_id=category.id,
                mentioned_brands=mentioned_brands,
                style_config=style_config,
                insert_config=insert_config
            )
            
            # 插入图片到文档
            if image_path and os.path.exists(image_path):
                paragraph = doc.add_paragraph()
                run = paragraph.add_run()
                
                # 使用配置的图片宽度
                image_width = style_config.get('image_width', 15) if style_config else 15
                run.add_picture(image_path, width=Inches(image_width / 2.54))  # 厘米转英寸
                
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                logger.info(f"✓ 对比表图片已插入: {image_path}")
            else:
                logger.warning(f"对比表图片生成失败或文件不存在: {image_path}")
                
        except Exception as e:
            logger.error(f"插入对比表图片失败: {e}")
            import traceback
            logger.error(traceback.format_exc())
    
    def _extract_mentioned_brands(self, text: str) -> List[str]:
        """
        从文本中提取提及的品牌（仅完整匹配）
        
        Args:
            text: 文档文本
        
        Returns:
            品牌名称列表
        """
        mentioned_brands = []
        
        # 获取所有品牌
        categories = self.comparison_db.get_all_categories()
        for category in categories:
            brands = self.comparison_db.get_brands_by_category(category.id)
            for brand in brands:
                brand_name = brand.name
                
                # 仅完整匹配（精确匹配完整品牌名）
                if brand_name in text:
                    mentioned_brands.append(brand_name)
                    logger.debug(f"品牌完整匹配: {brand_name}")
        
        logger.info(f"识别到的品牌: {mentioned_brands if mentioned_brands else '无'}")
        return mentioned_brands
    
    def _load_comparison_config(self) -> Optional[Dict]:
        """
        加载对比表配置
        
        Returns:
            配置字典
        """
        if not COMPARISON_TABLE_AVAILABLE:
            return None
            
        try:
            insert_config = self.comparison_db.get_config('insert_strategy')
            style_config = self.comparison_db.get_config('table_style')
            
            if insert_config or style_config:
                return {
                    'insert_strategy': insert_config,
                    'table_style': style_config
                }
            return None
        except Exception as e:
            logger.warning(f"加载对比表配置失败: {e}")
            return None
    
    def _add_heading(self, doc: Document, text: str, level: int):
        """
        添加标题
        
        Args:
            doc: Document 对象
            text: 文本内容
            level: 标题级别（1-3）
        """
        heading = doc.add_heading(text, level=level)
        self._apply_chinese_font(heading)
    
    def _add_paragraph(self, doc: Document, text: str):
        """
        添加段落
        
        Args:
            doc: Document 对象
            text: 文本内容
        """
        paragraph = doc.add_paragraph(text)
        self._apply_chinese_font(paragraph)
        self._apply_bold_keywords(paragraph)
        
        # 设置行间距和段后间距
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_after = Pt(10)
    
    def _add_list_item(self, doc: Document, text: str):
        """
        添加列表项
        
        Args:
            doc: Document 对象
            text: 文本内容
        """
        paragraph = doc.add_paragraph(text, style='List Bullet')
        self._apply_chinese_font(paragraph)
    
    def _add_image(self, doc: Document, image_folder: str, index: int):
        """
        添加图片
        
        Args:
            doc: Document 对象
            image_folder: 图片文件夹路径
            index: 索引（用于循环选择）
        """
        try:
            # 获取图片列表
            image_files = FileHandler.get_image_files(image_folder)
            if not image_files:
                logger.warning(f"图片文件夹为空: {image_folder}")
                return
            
            # 循环获取图片
            image_path = self.image_processor.get_cyclic_image(image_files, index)
            if not image_path:
                return
            
            # 获取 Alt 文本
            alt_text = self.image_processor.get_image_alt_text(image_path)
            
            # 插入图片
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            picture = run.add_picture(image_path, width=Inches(5))
            
            # 设置 Alt 文本
            picture._element.attrib['descr'] = alt_text
            
            # 居中对齐
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            logger.debug(f"插入图片: {alt_text}")
            
        except Exception as e:
            logger.error(f"插入图片失败: {e}")
    
    def _apply_chinese_font(self, paragraph):
        """
        应用中文字体（修复 python-docx 中文显示问题）
        
        Args:
            paragraph: 段落对象
        """
        for run in paragraph.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    def _apply_bold_keywords(self, paragraph):
        """
        对关键词加粗
        
        Args:
            paragraph: 段落对象
        """
        keywords = self.config.bold_keywords
        if not keywords:
            return
        
        for run in paragraph.runs:
            for keyword in keywords:
                if keyword in run.text:
                    run.font.bold = True
                    break
    
    def _generate_filename(self, row_data: List[str], index: int) -> str:
        """
        生成文件名
        
        Args:
            row_data: 行数据
            index: 索引
            
        Returns:
            文件名
        """
        # 尝试使用第一列（通常是标题）作为文件名
        if row_data and row_data[0]:
            title = row_data[0][:30]  # 限制长度
            # 清理文件名
            from ..utils.validators import sanitize_filename
            filename = sanitize_filename(title)
        else:
            filename = f"document_{index + 1}"
        
        return f"{filename}.docx"
    
    def _transpose_grid(self, grid_data: List[List[str]]) -> List[List[str]]:
        """
        转置网格数据（行列互换）
        
        Args:
            grid_data: 网格数据
            
        Returns:
            转置后的数据
        """
        if not grid_data:
            return []
        
        max_cols = max(len(row) for row in grid_data)
        columns = []
        
        for col_idx in range(max_cols):
            column = []
            for row in grid_data:
                if col_idx < len(row):
                    column.append(row[col_idx])
                else:
                    column.append("")
            columns.append(column)
        
        return columns
    
    def _extract_full_text(self, row_data: List[str]) -> str:
        """
        提取一行数据的全文本（用于查重）
        
        Args:
            row_data: 行数据
            
        Returns:
            合并后的文本
        """
        full_text = []
        
        for col_idx, cell_content in enumerate(row_data):
            if not cell_content or not cell_content.strip():
                continue
            
            # 获取列类型
            col_type = self.config.get_column_type(col_idx)
            
            # 忽略此列
            if col_type == 'Ignore':
                continue
            
            # 解析 Spintax
            parsed_content = self.spintax_parser.parse(cell_content)
            full_text.append(parsed_content)
        
        return " ".join(full_text)


if __name__ == "__main__":
    # 测试代码
    from ..utils.logger import setup_logger
    from ..config.settings import create_default_config
    
    setup_logger()
    
    # 创建测试配置
    config = create_default_config()
    config.set_column_type(0, 'H1', '标题')
    config.set_column_type(1, 'Body', '内容')
    
    # 创建生成器
    generator = DocumentGenerator(config)
    
    # 测试数据
    test_data = [
        ["测试标题1", "这是{第一|首个}测试内容"],
        ["测试标题2", "这是{第二|另一个}测试内容"],
    ]
    
    # 按行生成
    files = generator.generate_by_row(test_data, "test_output")
    logger.info(f"生成的文件: {files}")

